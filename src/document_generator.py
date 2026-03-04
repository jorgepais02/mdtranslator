"""Generate academic-style DOCX documents from Markdown files.

Converts a Markdown file into a DOCX document with:
  - Academic typography (Times New Roman, CJK, Arabic fonts)
  - Right-to-left support for Arabic/Hebrew
  - Header image from public/header.png
  - Page numbering in footer
  - Structured headings, bullet lists, numbered lists

Usage:
    python document_generator.py apuntes.md [output.docx] [--lang es]
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm


# ═══════════════════════════════════════════
# Constants & Configuration
# ═══════════════════════════════════════════

PROJECT_ROOT = Path(__file__).resolve().parent.parent
# Load config for header image
import json

def _load_config() -> dict:
    config_path = PROJECT_ROOT / "config.json"
    if not config_path.exists():
        config_path = PROJECT_ROOT / "config.example.json"
    if config_path.exists():
        try:
            return json.loads(config_path.read_text(encoding="utf-8"))
        except Exception:
            return {}
    return {}

CONFIG = _load_config()
header_cfg = CONFIG.get("document", {}).get("header_image")
HEADER_IMAGE = (PROJECT_ROOT / header_cfg) if header_cfg else None

# Languages that use right-to-left script
RTL_LANGS = {"ar", "he", "fa", "ur"}

# Font selection by language
FONT_MAP = {
    "zh": "SimSun",  # Chinese – common CJK serif
    "ar": "Amiri",  # Arabic – elegant Naskh (user requested)
}
DEFAULT_FONT = "Times New Roman"
MONOSPACE_FONT = "Courier New"

# Inline formatting regex (order matters: bold before italic)
_INLINE_RE = re.compile(
    r'(?P<bold_italic>\*\*\*(.+?)\*\*\*)'
    r'|(?P<bold>\*\*(.+?)\*\*)'
    r'|(?P<italic>\*(.+?)\*)'
    r'|(?P<code>`([^`]+)`)'
    r'|(?P<link>\[([^\]]+)\]\(([^)]+)\))'
)


# ═══════════════════════════════════════════
# Typography helpers
# ═══════════════════════════════════════════

def _resolve_font(lang: str) -> str:
    """Return the appropriate font name for the given language."""
    return FONT_MAP.get(lang, DEFAULT_FONT)


def _is_rtl(lang: str) -> bool:
    """Check whether the language uses right-to-left script."""
    return lang in RTL_LANGS


# ═══════════════════════════════════════════
# Document layout
# ═══════════════════════════════════════════

def configure_academic_margins(doc: Document) -> None:
    """Apply standard academic margins (2.5 cm all sides)."""
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def configure_base_typography(doc: Document, lang: str = "es") -> None:
    """Configure base document typography, adapting for language."""
    font_name = _resolve_font(lang)

    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(12)
    normal.font.bold = False
    normal.font.color.rgb = RGBColor(0, 0, 0)

    h1 = doc.styles["Heading 1"]
    h1.font.name = font_name
    h1.font.size = Pt(20)
    h1.font.bold = False
    h1.font.color.rgb = RGBColor(0, 0, 0)

    h2 = doc.styles["Heading 2"]
    h2.font.name = font_name
    h2.font.size = Pt(16)
    h2.font.bold = False
    h2.font.color.rgb = RGBColor(0, 0, 0)

    # Set Complex Script (CS) font and BiDi status for RTL languages
    if _is_rtl(lang):
        for style_name in ["Normal", "Heading 1", "Heading 2"]:
            style = doc.styles[style_name]
            rPr = style._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn("w:cs"), font_name)
            
            rtl = OxmlElement("w:rtl")
            rtl.set(qn("w:val"), "1")
            rPr.append(rtl)
            
            pPr = style._element.get_or_add_pPr()
            bidi = OxmlElement("w:bidi")
            bidi.set(qn("w:val"), "1")
            pPr.append(bidi)


# ═══════════════════════════════════════════
# RTL helpers
# ═══════════════════════════════════════════

def _set_paragraph_rtl(p) -> None:
    """Mark a paragraph as right-to-left at the XML level."""
    pPr = p._p.get_or_add_pPr()
    # Check if bidi is already there
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        pPr.insert(0, bidi)
    else:
        bidi.set(qn("w:val"), "1")
    
    # Remove any explicit jc (justification) — bidi=1 naturally right-aligns
    # Setting jc=right with bidi confuses some renderers (Google Docs import, Pages)
    jc = pPr.find(qn("w:jc"))
    if jc is not None:
        pPr.remove(jc)


def _set_run_rtl(run, font_name: str | None = None) -> None:
    """Mark a run as right-to-left and set complex-script font."""
    rPr = run._r.get_or_add_rPr()
    rtl_elem = rPr.find(qn("w:rtl"))
    if rtl_elem is None:
        rtl_elem = OxmlElement("w:rtl")
        rtl_elem.set(qn("w:val"), "1")
        rPr.append(rtl_elem)
    else:
        rtl_elem.set(qn("w:val"), "1")
    
    if font_name:
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:cs"), font_name)
        # For mixed text, setting ASCII and HAnsi helps Word stay stable
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")


def _apply_rtl_if_needed(p, lang: str) -> None:
    """Apply RTL paragraph direction and run flags when required."""
    if not _is_rtl(lang):
        return
    _set_paragraph_rtl(p)
    # Do NOT set p.alignment = RIGHT here — bidi handles alignment naturally.
    # Setting jc=right with bidi can confuse renderers.
    font_name = _resolve_font(lang)
    for run in p.runs:
        _set_run_rtl(run, font_name)


# ═══════════════════════════════════════════
# Spacing constants
# ═══════════════════════════════════════════

BODY_LINE_SPACING = 1.2
BODY_SPACE_AFTER_PT = 10
TITLE_SPACE_AFTER_PT = 36

H1_SPACE_BEFORE_PT = 28
H1_SPACE_AFTER_PT = 6
H2_SPACE_BEFORE_PT = 20
H2_SPACE_AFTER_PT = 5

SPACE_AFTER_BEFORE_HEADING_PT = 2

LIST_ITEM_SPACE_BEFORE_PT = 0
LIST_ITEM_SPACE_AFTER_PT = 6
LIST_BLOCK_END_SPACE_AFTER_PT = 12

LIST_BASE_LEFT_INDENT_CM = 1.5
LIST_LEVEL_INDENT_CM = 0.5
LIST_HANGING_CM = 0.4

BULLET_INDENT_SPACES = 2


# ═══════════════════════════════════════════
# Paragraph formatting
# ═══════════════════════════════════════════

def format_body_paragraph(p, lang: str = "es") -> None:
    """Apply body-text formatting to a paragraph."""
    font_name = _resolve_font(lang)
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if not _is_rtl(lang) else None
    # For RTL, bidi handles alignment naturally — don't set alignment explicitly
    
    line_spacing = 1.3 if _is_rtl(lang) else 1.2
    space_after = Pt(8) if _is_rtl(lang) else Pt(6)
    
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = space_after

    for run in p.runs:
        run.font.name = font_name
        run.font.size = Pt(12)
        run.font.bold = False
        run.font.color.rgb = RGBColor(0, 0, 0)

    _apply_rtl_if_needed(p, lang)


def format_heading_paragraph(p, level: int, lang: str = "es") -> None:
    """Apply heading formatting to a paragraph at the given level."""
    if not _is_rtl(lang):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # For RTL, bidi handles alignment

    if level == 1:
        p.paragraph_format.space_before = Pt(H1_SPACE_BEFORE_PT)
        p.paragraph_format.space_after = Pt(H1_SPACE_AFTER_PT)
    elif level == 2:
        p.paragraph_format.space_before = Pt(H2_SPACE_BEFORE_PT)
        p.paragraph_format.space_after = Pt(H2_SPACE_AFTER_PT)
    else:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)

    for run in p.runs:
        run.font.bold = False
        run.font.color.rgb = RGBColor(0, 0, 0)

    _apply_rtl_if_needed(p, lang)


def reduce_spacing_before_heading(doc: Document, pt: int = SPACE_AFTER_BEFORE_HEADING_PT) -> None:
    """Reduce spacing on the last Normal paragraph before a heading."""
    if not doc.paragraphs:
        return
    last = doc.paragraphs[-1]
    style_name = getattr(getattr(last, "style", None), "name", "") or ""
    if style_name == "Normal":
        last.paragraph_format.space_after = Pt(pt)


def reduce_spacing_before_list(doc: Document, pt: int = 0) -> None:
    """Reduce spacing on the last Normal paragraph before a list block."""
    if not doc.paragraphs:
        return
    last = doc.paragraphs[-1]
    style_name = getattr(getattr(last, "style", None), "name", "") or ""
    if style_name == "Normal":
        last.paragraph_format.space_after = Pt(pt)


def ensure_spacing_after_list(doc: Document) -> None:
    """Add extra spacing after the last list item if needed."""
    if not doc.paragraphs:
        return
    last = doc.paragraphs[-1]
    style_name = getattr(getattr(last, "style", None), "name", "") or ""
    if "List" in style_name:
        last.paragraph_format.space_after = Pt(LIST_BLOCK_END_SPACE_AFTER_PT)


# ═══════════════════════════════════════════
# Header & Footer
# ═══════════════════════════════════════════

def insert_header_image(doc: Document) -> None:
    """Insert header.png spanning the full page width (edge to edge)."""
    if not HEADER_IMAGE or not HEADER_IMAGE.exists():
        return

    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Remove paragraph spacing so the image sits flush
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # Negative indents to compensate for page margins → true edge-to-edge
    left_margin = section.left_margin
    right_margin = section.right_margin
    p.paragraph_format.left_indent = -left_margin
    p.paragraph_format.right_indent = -right_margin

    # Clear existing content
    for run in p.runs:
        run.clear()

    run = p.add_run()
    # Full A4 page width: 21 cm
    page_width = section.page_width
    run.add_picture(str(HEADER_IMAGE), width=page_width)


def insert_page_number_footer(doc: Document) -> None:
    """Insert an automatic page number in the footer (right-aligned)."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = p.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(fld_end)

    run.font.name = DEFAULT_FONT
    run.font.size = Pt(12)
    run.font.bold = False
    run.font.color.rgb = RGBColor(0, 0, 0)


# ═══════════════════════════════════════════
# Content insertion helpers
# ═══════════════════════════════════════════

def insert_document_title(doc: Document, text: str, lang: str = "es") -> None:
    """Insert the main document title (centered, large font)."""
    font_name = _resolve_font(lang)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(text.strip())
    run.bold = False
    run.font.name = font_name
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 0, 0)

    p.paragraph_format.space_after = Pt(TITLE_SPACE_AFTER_PT)
    _apply_rtl_if_needed(p, lang)


def insert_heading(doc: Document, text: str, md_level: int, lang: str = "es") -> None:
    """Insert a section heading (maps Markdown level to Word level)."""
    word_level = max(1, md_level - 1)
    word_level = min(word_level, 9)
    p = doc.add_heading(text.strip(), level=word_level)
    format_heading_paragraph(p, level=word_level, lang=lang)


def insert_body_paragraph(doc: Document, text: str, lang: str = "es") -> None:
    """Insert a body paragraph with standard formatting."""
    p = doc.add_paragraph(text.rstrip())
    format_body_paragraph(p, lang=lang)


# ═══════════════════════════════════════════
# Inline formatting helpers
# ═══════════════════════════════════════════

def _parse_inline_formatting(text: str) -> list[dict]:
    """Split text into segments with inline formatting metadata.

    Each segment is a dict with keys:
      text, bold, italic, code, link_url
    """
    segments: list[dict] = []
    last_end = 0

    for m in _INLINE_RE.finditer(text):
        # Plain text before this match
        if m.start() > last_end:
            segments.append({
                "text": text[last_end:m.start()],
                "bold": False, "italic": False, "code": False, "link_url": None,
            })

        if m.group("bold_italic"):
            segments.append({
                "text": m.group(2),
                "bold": True, "italic": True, "code": False, "link_url": None,
            })
        elif m.group("bold"):
            segments.append({
                "text": m.group(4),
                "bold": True, "italic": False, "code": False, "link_url": None,
            })
        elif m.group("italic"):
            segments.append({
                "text": m.group(6),
                "bold": False, "italic": True, "code": False, "link_url": None,
            })
        elif m.group("code"):
            segments.append({
                "text": m.group(8),
                "bold": False, "italic": False, "code": True, "link_url": None,
            })
        elif m.group("link"):
            segments.append({
                "text": m.group(10),
                "bold": False, "italic": False, "code": False, "link_url": m.group(11),
            })
        last_end = m.end()

    # Trailing plain text
    if last_end < len(text):
        segments.append({
            "text": text[last_end:],
            "bold": False, "italic": False, "code": False, "link_url": None,
        })

    # If no formatting found, return single plain segment
    if not segments:
        segments.append({
            "text": text,
            "bold": False, "italic": False, "code": False, "link_url": None,
        })

    return segments


def _add_formatted_runs(p, text: str, lang: str = "es", font_size_pt: int = 12) -> None:
    """Add runs with inline formatting (bold, italic, code, links) to a paragraph."""
    font_name = _resolve_font(lang)
    segments = _parse_inline_formatting(text.strip())

    for seg in segments:
        run = p.add_run(seg["text"])
        run.font.name = font_name if not seg["code"] else MONOSPACE_FONT
        run.font.size = Pt(font_size_pt if not seg["code"] else max(font_size_pt - 1, 9))
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = seg["bold"]
        run.italic = seg["italic"]

        if seg["code"]:
            # Light grey background shading for inline code
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "E8E8E8")
            run._r.get_or_add_rPr().append(shd)

        if seg["link_url"]:
            run.font.color.rgb = RGBColor(0, 0, 238)
            run.font.underline = True


def _apply_list_indentation(p, level: int, lang: str = "es") -> None:
    """Set list indentation, respecting RTL for Arabic-script languages."""
    if _is_rtl(lang):
        # RTL: Word behavior for 'left' indent is actually 'start' if bidi is on? No, it's literal.
        # We must set both left and right carefully.
        p.paragraph_format.left_indent = Cm(0)  # Reset left
        right = LIST_BASE_LEFT_INDENT_CM + (level * LIST_LEVEL_INDENT_CM)
        p.paragraph_format.right_indent = Cm(right)
        p.paragraph_format.first_line_indent = Cm(-LIST_HANGING_CM)
    else:
        p.paragraph_format.right_indent = Cm(0) # Reset right
        left = LIST_BASE_LEFT_INDENT_CM + (level * LIST_LEVEL_INDENT_CM)
        p.paragraph_format.left_indent = Cm(left)
        p.paragraph_format.first_line_indent = Cm(-LIST_HANGING_CM)


def insert_bullet_item(doc: Document, text: str, level: int = 0, lang: str = "es") -> None:
    """Insert a bullet list item at the given nesting level."""
    if _is_rtl(lang):
        # Use Normal style for RTL to avoid LTR-hardcoded bullet behaviors
        p = doc.add_paragraph()
        _set_paragraph_rtl(p)
        # bidi handles alignment naturally
        # Manual bullet character for RTL - using Arabic bullet '\u2022' or similar
        text = f"• {text}"
    else:
        style = "List Bullet" if level == 0 else "List Bullet 2" if level == 1 else "List Bullet 3"
        p = doc.add_paragraph(style=style)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.3 if _is_rtl(lang) else 1.2
    p.paragraph_format.space_before = Pt(LIST_ITEM_SPACE_BEFORE_PT)
    p.paragraph_format.space_after = Pt(LIST_ITEM_SPACE_AFTER_PT)

    _apply_list_indentation(p, level, lang=lang)
    _add_formatted_runs(p, text, lang=lang)
    _apply_rtl_if_needed(p, lang)


def insert_numbered_item(doc: Document, text: str, level: int = 0, lang: str = "es") -> None:
    """Insert a numbered list item at the given nesting level."""
    if _is_rtl(lang):
        p = doc.add_paragraph()
        # bidi handles alignment naturally via _apply_rtl_if_needed below
    else:
        style = "List Number" if level == 0 else "List Number 2" if level == 1 else "List Number 3"
        p = doc.add_paragraph(style=style)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.3 if _is_rtl(lang) else 1.2
    p.paragraph_format.space_before = Pt(LIST_ITEM_SPACE_BEFORE_PT)
    p.paragraph_format.space_after = Pt(LIST_ITEM_SPACE_AFTER_PT)

    _apply_list_indentation(p, level, lang=lang)
    _add_formatted_runs(p, text, lang=lang)
    _apply_rtl_if_needed(p, lang)


def insert_alphabetic_item(doc: Document, label: str, text: str, level: int = 0, lang: str = "es") -> None:
    """Insert an alphabetic list item (e.g., A) or b.) with manual labels."""
    # We use Normal style but apply list-like indentation
    p = doc.add_paragraph()
    if not _is_rtl(lang):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.3 if _is_rtl(lang) else 1.2
    p.paragraph_format.space_before = Pt(LIST_ITEM_SPACE_BEFORE_PT)
    p.paragraph_format.space_after = Pt(LIST_ITEM_SPACE_AFTER_PT)

    _apply_list_indentation(p, level, lang=lang)
    
    # Prepend the label (e.g., "A) ") to the text so it matches the input
    full_text = f"{label} {text}"
    _add_formatted_runs(p, full_text, lang=lang)
    _apply_rtl_if_needed(p, lang)


# ═══════════════════════════════════════════
# Blockquote & Code Block helpers
# ═══════════════════════════════════════════

def insert_blockquote(doc: Document, text: str, lang: str = "es") -> None:
    """Insert a blockquote paragraph with left border and indent."""
    font_name = _resolve_font(lang)
    p = doc.add_paragraph()
    if not _is_rtl(lang):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # For RTL, bidi from _apply_rtl_if_needed handles alignment
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2

    # Left border via XML
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left_bdr = OxmlElement("w:left")
    left_bdr.set(qn("w:val"), "single")
    left_bdr.set(qn("w:sz"), "12")
    left_bdr.set(qn("w:space"), "8")
    left_bdr.set(qn("w:color"), "999999")
    pBdr.append(left_bdr)
    pPr.append(pBdr)

    # Paragraph shading
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)

    _add_formatted_runs(p, text, lang=lang)
    _apply_rtl_if_needed(p, lang)


def insert_code_block_line(doc: Document, text: str, lang: str = "es") -> None:
    """Insert a single line of a fenced code block."""
    p = doc.add_paragraph(text)
    format_body_paragraph(p, lang=lang)


def insert_simple_table(doc: Document, header_row: list[str], data_rows: list[list[str]], lang: str = "es") -> None:
    """Insert a simple markdown table into the document."""
    font_name = _resolve_font(lang)
    cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.style = "Table Grid"

    # Header
    for i, h in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h.strip())
        run.bold = True
        run.font.name = font_name
        run.font.size = Pt(11)

    # Data
    for r_idx, row in enumerate(data_rows):
        for c_idx, val in enumerate(row):
            if c_idx >= cols:
                break
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val.strip())
            run.font.name = font_name
            run.font.size = Pt(11)


# ═══════════════════════════════════════════
# Markdown → DOCX conversion
# ═══════════════════════════════════════════

HEADING_RE = re.compile(r"^(#{1,6})\s+(.*\S)\s*$")
BLOCKQUOTE_RE = re.compile(r"^>\s?(.*)") 
CODE_FENCE_RE = re.compile(r"^```")
TABLE_ROW_RE = re.compile(r"^\|(.+)\|\s*$")
TABLE_SEP_RE = re.compile(r"^\|[\s\-:|]+\|\s*$")


def convert_markdown_to_docx(md_path: Path, docx_path: Path, lang: str = "es") -> None:
    """Convert a Markdown file into a formatted DOCX document."""
    md_lines = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document()
    configure_academic_margins(doc)
    configure_base_typography(doc, lang=lang)
    insert_header_image(doc)
    insert_page_number_footer(doc)

    # Set document-level RTL if needed
    if _is_rtl(lang):
        for section in doc.sections:
            sectPr = section._sectPr
            bidi = OxmlElement("w:bidi")
            bidi.set(qn("w:val"), "1")
            sectPr.append(bidi)
            rtl_gutter = OxmlElement("w:rtlGutter")
            rtl_gutter.set(qn("w:val"), "1")
            sectPr.append(rtl_gutter)

    title_written = False
    last_was_heading = False  # Track consecutive headings for spacing fix
    last_element_type = 'NORMAL'
    last_list_level = 0
    buffer: list[str] = []
    in_code_block = False
    code_block_lines: list[str] = []
    table_rows: list[str] = []  # Accumulate raw table rows

    def flush_buffer() -> None:
        nonlocal buffer
        para = " ".join(line.strip() for line in buffer).strip()
        buffer_type = last_element_type
        buffer = []
        if para:
            ensure_spacing_after_list(doc)
            p = doc.add_paragraph()
            _add_formatted_runs(p, para, lang=lang)
            if buffer_type == 'BULLET_CONT':
                p.paragraph_format.line_spacing = 1.3 if _is_rtl(lang) else 1.2
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(LIST_ITEM_SPACE_AFTER_PT)
                _apply_list_indentation(p, last_list_level, lang=lang)
                _apply_rtl_if_needed(p, lang)
            else:
                format_body_paragraph(p, lang=lang)

    def flush_code_block() -> None:
        nonlocal code_block_lines
        for cl in code_block_lines:
            insert_code_block_line(doc, cl, lang=lang)
        code_block_lines = []

    def flush_table() -> None:
        nonlocal table_rows
        if len(table_rows) < 2:
            # Not a real table, treat as body text
            for row in table_rows:
                buffer.append(row)
            table_rows = []
            return
        # Parse header and data
        header_cells = [c.strip() for c in table_rows[0].strip("|").split("|")]
        data: list[list[str]] = []
        for row in table_rows[1:]:
            if TABLE_SEP_RE.match(row):
                continue  # skip separator row
            cells = [c.strip() for c in row.strip("|").split("|")]
            data.append(cells)
        if data:
            insert_simple_table(doc, header_cells, data, lang=lang)
        table_rows = []

    for line in md_lines:
        # ── Code blocks ──────────────────────────
        if CODE_FENCE_RE.match(line.strip()):
            if in_code_block:
                flush_code_block()
                in_code_block = False
            else:
                flush_buffer()
                in_code_block = True
            continue
        if in_code_block:
            code_block_lines.append(line)
            continue

        # ── Table accumulation ───────────────────
        if TABLE_ROW_RE.match(line):
            if not table_rows:
                flush_buffer()
            table_rows.append(line)
            last_was_heading = False
            continue
        elif table_rows:
            flush_table()

        # ── Headings ─────────────────────────────
        m = HEADING_RE.match(line)
        if m:
            flush_buffer()
            hashes, heading_text = m.group(1), m.group(2)
            level = len(hashes)

            if level == 1 and not title_written:
                insert_document_title(doc, heading_text, lang=lang)
                title_written = True
            else:
                if last_was_heading:
                    # Consecutive headings → reduce gap between them
                    prev_p = doc.paragraphs[-1]
                    prev_p.paragraph_format.space_after = Pt(4)
                else:
                    reduce_spacing_before_heading(doc)
                insert_heading(doc, heading_text, md_level=level, lang=lang)
                if last_was_heading:
                    # Also reduce space_before on this new heading
                    doc.paragraphs[-1].paragraph_format.space_before = Pt(6)
            last_was_heading = True
            continue

        last_was_heading = False

        # ── Blank lines ──────────────────────────
        if line.strip() == "":
            flush_buffer()
            last_element_type = 'NORMAL'
            continue

        # ── Blockquotes ──────────────────────────
        bq = BLOCKQUOTE_RE.match(line)
        if bq:
            flush_buffer()
            insert_blockquote(doc, bq.group(1), lang=lang)
            continue

        # ── Horizontal rule ──────────────────────
        if re.match(r"^\s*(---|\*\*\*)\s*$", line):
            flush_buffer()
            buffer.append(line)
            continue

        # ── LIST DETECTION ───────────────────────
        raw = line.rstrip("\n")
        leading = len(raw) - len(raw.lstrip(" \t"))
        leading_spaces = raw[:leading].replace("\t", " " * 4)
        indent = len(leading_spaces)
        level = indent // BULLET_INDENT_SPACES
        level = min(level, 2)

        s = raw.lstrip(" \t").strip()

        # Numbered list
        m_num = re.match(r"^(\d+)\.\s+(.*\S)\s*$", s)
        if m_num:
            flush_buffer()
            reduce_spacing_before_list(doc, pt=0)
            content = m_num.group(2).strip()
            insert_numbered_item(doc, content, level=level, lang=lang)
            last_element_type = 'LIST'
            last_list_level = level
            continue

        # Alphabetic list (e.g., A) or b.)
        m_alpha = re.match(r"^([a-zA-Z])([\)\.])\s+(.*\S)\s*$", s)
        if m_alpha:
            flush_buffer()
            reduce_spacing_before_list(doc, pt=0)
            label = m_alpha.group(1) + m_alpha.group(2)
            content = m_alpha.group(3).strip()
            insert_alphabetic_item(doc, label, content, level=level, lang=lang)
            last_element_type = 'LIST'
            last_list_level = level
            continue

        # Bullet list
        if s.startswith("- "):
            flush_buffer()
            reduce_spacing_before_list(doc, pt=0)
            content = s[2:].strip()

            if " - " in content:
                parts = [p.strip() for p in content.split(" - ") if p.strip()]
                for part in parts:
                    if part.startswith("- "):
                        part = part[2:].strip()
                    insert_bullet_item(doc, part, level=level, lang=lang)
            else:
                insert_bullet_item(doc, content, level=level, lang=lang)
            last_element_type = 'LIST'
            last_list_level = level
            continue

        # Check for list continuation (indented paragraph)
        if last_element_type in ('LIST', 'BULLET_CONT') and (line.startswith("  ") or line.startswith("\t")):
            if last_element_type == 'LIST':
                # Starting a new continuation paragraph for the current list item
                last_element_type = 'BULLET_CONT'
            buffer.append(line)
            continue
            
        if last_element_type == 'BULLET_CONT':
            flush_buffer()

        last_element_type = 'NORMAL'
        buffer.append(line)

    # Final flushes
    if in_code_block:
        flush_code_block()
    if table_rows:
        flush_table()
    flush_buffer()
    doc.save(docx_path)


# ═══════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════

def main() -> int:
    parser = argparse.ArgumentParser(description="Convert Markdown to DOCX with academic formatting.")
    parser.add_argument("input", type=Path, help="Path to the input Markdown file.")
    parser.add_argument("output", type=Path, nargs="?", help="Optional path for the output DOCX. Defaults to input with .docx extension.")
    parser.add_argument("--lang", default="es", help="Language code for formatting (e.g., 'es', 'ar', 'zh'). Default is 'es'.")
    
    # Check if we are being called via legacy sys.argv style or let argparse handle sys.argv
    # To keep it simple and standard:
    args = parser.parse_args()

    md_path = args.input.expanduser().resolve()
    if not md_path.exists():
        print(f"ERROR: File not found: {md_path}")
        return 2

    # If output is not provided, use input name but with .docx
    out = args.output if args.output else md_path.with_suffix(".docx")
    out = out.expanduser().resolve()
    
    lang = args.lang.lower()

    convert_markdown_to_docx(md_path, out, lang=lang)

    print(f"OK -> {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
